import os
from docx import Document

class WordSkill:
    def __init__(self):
        pass

    def read_docx(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            return f"Error: File {file_path} does not exist."
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            return f"Error reading file {file_path}: {e}"

    def write_docx(self, file_path: str, text: str) -> str:
        try:
            doc = Document()
            for line in text.split('\n'):
                doc.add_paragraph(line)
            doc.save(file_path)
            return f"Successfully wrote to {file_path}" 
        except Exception as e:
            return f"Error writing file {file_path}: {e}"

# docファイル対応は別途検討とする


if __name__ == '__main__':
    ws = WordSkill()
    print(ws.read_docx('sample.docx'))
    print(ws.write_docx('output.docx', 'Hello\nWorld'))
    
